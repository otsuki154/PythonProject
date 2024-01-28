import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, RGBColor


def addTextWithFont(text, name, size, rgb_values):
    pageTitle = doc.add_paragraph()
    pageRun = pageTitle.add_run(text)
    pageFont = pageRun.font
    pageFont.name = name
    pageFont.size = Pt(size)
    pageFont.color.rgb = RGBColor(*rgb_values)


# URL của trang web
# url = "https://vnexpress.net/the-thao"
url = "https://vnexpress.net/khoa-hoc"

# Gửi yêu cầu GET đến trang web và lấy nội dung HTML
response = requests.get(url)
html_content = response.text

# Sử dụng BeautifulSoup để phân tích HTML
soup = BeautifulSoup(html_content, "html.parser")

# Lấy danh sách các bài báo
articles = soup.select("p.description a")

# Tạo và lưu vào file Word
doc = Document()
addTextWithFont("Top 10 Bài Báo Khoa Học Mới Nhất", 'Courier New', 20, [0, 128, 128])

for index, article in enumerate(articles[:10], 1):
    if 'title' in article.attrs:
        #lấy nội dung title
        title = article['title']
    else:
        continue # nếu không có attrs title thì bỏ qua luôn bài báo

    link = article["href"] #lấy link bài báo
    content = article.get_text(strip=True) #lấy nội dung tóm tắt

    addTextWithFont(f"{title}", 'Courier New', 16, [238, 3, 173]) #Thêm nội dung title vào file docx
    addTextWithFont(f"{content}", 'Courier New', 14, [122, 203, 173]) #Thêm nội dung contents vào file docx
    addTextWithFont(f"Link: {link}", 'Courier New', 12, [122, 255, 255]) #Thêm nội dung link vào file docx
    doc.add_paragraph("---------------------------------------------------")

# Lưu file Word
doc.save("khoahoc.docx")

print("File 'khoahoc.docx' đã được tạo.")
